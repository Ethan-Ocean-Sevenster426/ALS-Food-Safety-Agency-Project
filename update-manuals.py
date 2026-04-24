"""
Append a 'Recent Updates - April 2026' section to every EPVS user manual,
documenting all changes made in this iteration. Idempotent: re-running
removes the existing section and re-inserts the latest version.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK

REPO = os.path.dirname(os.path.abspath(__file__))

MANUALS = [
    "EPVS Super Admin User Manual.docx",
    "EPVS Admin User Manual.docx",
    "EPVS Inspector User Manual.docx",
    "EPVS Company Admin User Manual.docx",
    "EPVS User Manual.docx",
    "EPVS Process Flow Documentation.docx",
]

SECTION_TITLE = "Recent Updates - April 2026"

# (heading_level, text). Heading level 0 = section title (Heading 1),
# 1 = subheading (Heading 2), 2 = bullet, 3 = paragraph.
COMMON_UPDATES = [
    (1, "EPV period now reflects the previous month"),
    (3, "Whenever an EPV is issued (manually from the Consolidated Master "
        "Facility Database or Company Overview, or automatically each month), "
        "the verification is created for the previous calendar month. Example: "
        "an EPV issued in March 2026 captures February 2026 data."),

    (1, "Automatic monthly EPV sending"),
    (3, "Once a facility has been placed on the EPV cycle (status "
        "'On EPV Cycle'), the system now automatically issues that facility a "
        "fresh EPV at the start of every month covering the previous month's "
        "data. The scheduler runs once a day; if the previous-month EPV is "
        "already present, nothing is sent. The email goes to all four facility "
        "addresses on file (primary, owner, accounts and manager). To take a "
        "facility off the cycle, clear its EPVCycleStatus on the Consolidated "
        "Master Facility Database page."),

    (1, "Resend EPV email button"),
    (3, "The Company Overview window now shows a 'Resend Email' button "
        "alongside every Pending EPV. Use it when a facility reports they "
        "never received the original notification - it re-sends the same EPV "
        "link to all facility email addresses on file. No new EPV record is "
        "created and the resend is logged in EmailSendLog."),

    (1, "Facility verified status now visible"),
    (3, "The Consolidated Master Facility Database page shows a 'Verified' "
        "badge once the facility's invited Company Admin has accepted their "
        "invitation and completed the onboarding wizard. A separate 'On EPV "
        "Cycle' tag indicates the facility is on the monthly auto-send "
        "schedule. The status cell stays empty until one of those two events "
        "applies."),

    (1, "EPV form (Eggs page) - new 'Eggs Exported' deduction"),
    (3, "Section C (Deductions) now includes 'Eggs Exported'. Like the other "
        "deductions, it reduces the theoretical closing stock and does NOT "
        "affect the egg levy (only Section D Sales drives the levy)."),

    (1, "EPV form (Eggs page) - purchase source detail and attachments"),
    (3, "When Graded Eggs Purchased or Ungraded Eggs Purchased is greater "
        "than zero, a 'Source & Supplier Detail' comment box opens under "
        "Section B together with a file uploader. Facility users and "
        "inspectors can describe where and from whom the purchase was made, "
        "and attach supporting documents (PDF, PNG, or JPG, up to 15 MB each). "
        "Uploaded files appear with a 'Remove' button while the EPV is still "
        "editable, and as inline links in the Review section."),

    (1, "EPV form (Pulp page) - 'Conversion Loss' line"),
    (3, "Sales & Deductions on the pulp page now has a 'Conversion Loss' row "
        "directly below 'Sold to Other Producers'. Conversion loss does NOT "
        "contribute to the pulp levy, but it is subtracted from closing pulp "
        "stock so the carry-forward stays balanced."),

    (1, "EPV form (Pulp page) - pulp purchase source detail and attachments"),
    (3, "When 'Pulp Purchased from Others' is greater than zero, a comment "
        "box and uploader appear (same 15 MB limit and PDF/PNG/JPG rules) so "
        "the source of the purchased pulp can be recorded and supporting "
        "documents attached."),

    (1, "Verifications can only be captured for completed months"),
    (3, "The 'Add Verification' modal in Company Overview now defaults to the "
        "previous month, and the month picker hides the current and any "
        "future months. The backend rejects current-month submissions with a "
        "clear error message: a month must be over before its EPV can be "
        "captured."),
]


def remove_existing_section(doc, title):
    """Remove the previously appended 'Recent Updates' section, if any."""
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = None
    for i, child in enumerate(children):
        if child.tag.endswith('}p'):
            text = ''.join(t.text or '' for t in child.iter() if t.tag.endswith('}t'))
            if text.strip() == title:
                start_idx = i
                break
    if start_idx is not None:
        for child in children[start_idx:]:
            body.remove(child)


def has_style(doc, name):
    try:
        _ = doc.styles[name]
        return True
    except KeyError:
        return False


def add_styled_heading(doc, text, level):
    """Add a heading using built-in styles when available, else a bold run."""
    style_name = f'Heading {level}'
    if has_style(doc, style_name):
        return doc.add_heading(text, level=level)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    run.font.color.rgb = RGBColor(0x14, 0x4E, 0x7A)
    return p


def add_bullet(doc, text):
    if has_style(doc, 'List Bullet'):
        return doc.add_paragraph(text, style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('  - ' + text)
    return p


def append_section(path):
    doc = Document(path)
    remove_existing_section(doc, SECTION_TITLE)

    # Page break before the new section
    last_para = doc.add_paragraph()
    last_para.add_run().add_break(WD_BREAK.PAGE)

    add_styled_heading(doc, SECTION_TITLE, 1)
    intro = doc.add_paragraph()
    intro.add_run(
        "This section summarises the changes shipped in April 2026. Earlier "
        "sections of this manual remain accurate; the items below override or "
        "extend them."
    )

    for level, text in COMMON_UPDATES:
        if level == 1:
            add_styled_heading(doc, text, 2)
        elif level == 2:
            add_bullet(doc, text)
        else:
            doc.add_paragraph(text)

    doc.save(path)


def main():
    missing = [m for m in MANUALS if not os.path.exists(os.path.join(REPO, m))]
    if missing:
        print("Missing manuals:", missing)
        sys.exit(1)

    for m in MANUALS:
        path = os.path.join(REPO, m)
        try:
            append_section(path)
            print("updated:", m)
        except Exception as e:
            print("FAILED:", m, "->", e)


if __name__ == "__main__":
    main()
