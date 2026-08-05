"""
Regenerate every EPVS user manual from role-specific templates.

Rewrites all seven manuals from scratch using proper Word structure
(headings, tables, callouts, TOC, footer) so they navigate and print
professionally. External-audience manuals use a friendly voice;
internal manuals stay operational.

Requires python-docx (`pip install python-docx`).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

REPO = os.path.dirname(os.path.abspath(__file__))

# ── Palette ────────────────────────────────────────────────────────────
FSA_TEAL      = RGBColor(0x0E, 0x7C, 0x7B)
FSA_TEAL_DARK = RGBColor(0x0A, 0x5C, 0x5B)
DARK_INK      = RGBColor(0x1F, 0x29, 0x37)
BODY_INK      = RGBColor(0x37, 0x41, 0x51)
MUTED_INK     = RGBColor(0x6B, 0x72, 0x80)
LIGHT_INK     = RGBColor(0x9C, 0xA3, 0xAF)
NOTE_BG       = 'DBEAFE'  # light blue
NOTE_FG       = RGBColor(0x1E, 0x40, 0xAF)
TIP_BG        = 'D1FAE5'  # light green
TIP_FG        = RGBColor(0x04, 0x78, 0x57)
WARN_BG       = 'FEF3C7'  # light amber
WARN_FG       = RGBColor(0x92, 0x40, 0x0E)
ACCENT_RED    = RGBColor(0xDC, 0x26, 0x26)
TABLE_HDR_BG  = '0E7C7B'
TABLE_HDR_FG  = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_BG      = 'F3F4F6'


# ── Low-level XML helpers ──────────────────────────────────────────────
def _set_cell_shading(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _set_para_shading(paragraph, hex_fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)


def _set_para_border(paragraph, side='left', color='0E7C7B', size='24'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), size)
    b.set(qn('w:space'), '8')
    b.set(qn('w:color'), color)
    pBdr.append(b)


def _add_page_number(paragraph):
    """Insert PAGE and NUMPAGES field codes."""
    for field, label in [('PAGE', ''), (' of ', None), ('NUMPAGES', '')]:
        if label is None:
            r = paragraph.add_run(field)
            r.font.size = Pt(9)
            r.font.color.rgb = MUTED_INK
            continue
        run = paragraph.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED_INK
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def _add_toc(paragraph):
    """Insert a Word TOC field. Reader must press F9 or right-click → Update Field."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = 'Right-click and choose "Update Field" to build the Table of Contents.'
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    r2 = paragraph.add_run()
    r2._r.append(placeholder)
    r3 = paragraph.add_run()
    r3._r.append(fldChar3)


# ── High-level building blocks ─────────────────────────────────────────
def _configure_styles(doc):
    """Set body defaults + heading colors so TOC + navigation work in Word."""
    normal = doc.styles['Normal']
    normal.font.name = 'Segoe UI'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY_INK

    for name, size, color, bold in [
        ('Heading 1', 22, FSA_TEAL, True),
        ('Heading 2', 16, FSA_TEAL_DARK, True),
        ('Heading 3', 12, DARK_INK, True),
        ('Title', 32, FSA_TEAL, True),
        ('Subtitle', 16, MUTED_INK, False),
    ]:
        try:
            s = doc.styles[name]
            s.font.name = 'Segoe UI'
            s.font.size = Pt(size)
            s.font.color.rgb = color
            s.font.bold = bold
        except KeyError:
            pass


def _footer(doc, manual_name):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(manual_name + '   |   EPVS · ALS Food Safety Agency')
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED_INK
    r.italic = True
    p.add_run('\t\t')
    _add_page_number(p)


def cover(doc, *, title, subtitle, audience, edition='April 2026'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run('EPVS')
    run.font.size = Pt(72)
    run.font.color.rgb = FSA_TEAL
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Egg Production Verification System')
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED_INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(title)
    run.font.size = Pt(30)
    run.font.color.rgb = DARK_INK
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED_INK
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    run = p.add_run('Audience: ' + audience)
    run.font.size = Pt(12)
    run.font.color.rgb = FSA_TEAL_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Edition ' + edition)
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED_INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run('ALS Food Safety Agency')
    run.font.size = Pt(11)
    run.font.color.rgb = FSA_TEAL
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('All amounts throughout EPVS are shown EXCLUSIVE of VAT.')
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT_RED
    run.bold = True

    page_break(doc)


def toc_page(doc):
    h1(doc, 'Table of Contents')
    para(doc, 'This document uses a Word Table of Contents. To fill it in, right-click below and choose "Update Field" (or press F9). All headings you see through the guide are included.')
    p = doc.add_paragraph()
    _add_toc(p)
    page_break(doc)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY_INK
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet') if 'List Bullet' in [s.name for s in doc.styles] else doc.add_paragraph()
    r = p.add_run(text if 'List Bullet' in [s.name for s in doc.styles] else '  •  ' + text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY_INK
    return p


def numbered(doc, n, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f'  {n}.  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = FSA_TEAL_DARK
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BODY_INK
    return p


RED_BG = 'FEE2E2'
def callout(doc, kind, text):
    """kind ∈ 'note' | 'tip' | 'warning' | 'red'"""
    bg = {'note': NOTE_BG, 'tip': TIP_BG, 'warning': WARN_BG, 'red': RED_BG}[kind]
    fg = {'note': NOTE_FG, 'tip': TIP_FG, 'warning': WARN_FG, 'red': ACCENT_RED}[kind]
    label = {'note': 'NOTE', 'tip': 'TIP', 'warning': 'IMPORTANT', 'red': 'EXCL. VAT'}[kind]
    p = doc.add_paragraph()
    _set_para_shading(p, bg)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(label + '  ')
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = fg
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_INK
    return p


def reference_table(doc, headers, rows):
    """A styled reference table with header shading and a subtle zebra."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D1D5DB')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        _set_cell_shading(cell, TABLE_HDR_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = TABLE_HDR_FG

    for row_idx, row in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        for i, val in enumerate(row):
            cell = cells[i]
            if row_idx % 2 == 1:
                _set_cell_shading(cell, ZEBRA_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            r.font.color.rgb = BODY_INK
    doc.add_paragraph()  # breathing room after the table
    return table


def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run('')
    r.add_break(WD_BREAK.PAGE)


def section_intro(doc, purpose):
    """A visually distinct 'in this section' block at the top of a section."""
    p = doc.add_paragraph()
    _set_para_border(p, 'left', '0E7C7B', '24')
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run('IN THIS SECTION  ')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = FSA_TEAL
    r2 = p.add_run(purpose)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_INK
    r2.italic = True


# ── Shared content blocks ──────────────────────────────────────────────
def support_and_glossary(doc, *, external=False):
    h1(doc, 'Getting help')
    section_intro(doc, 'Where to reach out when something isn\'t working, and where to look for answers on your own.')
    h2(doc, 'Inside EPVS')
    para(doc,
         'Every page has a Support link in the top ribbon and a floating help '
         'button in the bottom-right corner. Click either to raise a ticket. Support '
         'tickets are triaged by category (Access, EPV Submission, Payment, Dashboard, '
         'Company & Facility, or General) and priority (Low, Medium, High, Urgent).')
    para(doc, 'You will get email updates as your ticket moves through Open → In Progress → Resolved. Follow-up comments can be added on the ticket at any time.')
    h2(doc, 'The User Manual button')
    para(doc,
         'The "User Manual" button in the top ribbon downloads this guide as a PDF at '
         'any time. Useful for offline reading, printing, or forwarding to a colleague.')
    if external:
        callout(doc, 'tip', 'If a support ticket is urgent (e.g. you cannot log in on the day an EPV is due), also email your account manager. EPVS notifications may take a few minutes to reach on-call staff after-hours.')

    h1(doc, 'Glossary')
    section_intro(doc, 'The vocabulary this manual uses.')
    reference_table(doc,
        headers=['Term', 'Meaning'],
        rows=[
            ['EPV', 'Egg Production Verification — the monthly declaration each facility submits.'],
            ['Levy', 'Statutory levy owed on eggs, pulp, and powder sold to trade.'],
            ['Facility', 'A registered producer, packer, or breaking-plant on the system.'],
            ['ALS', 'The 3rd-party levy collector (previously called "Super" in older releases).'],
            ['Inspector', 'ALS Food Safety Agency staff who verify facility EPVs.'],
            ['Reconciled', 'A payment has been received and matched to an EPV.'],
            ['POP', 'Proof of Payment — the bank confirmation the facility uploads after paying.'],
            ['On EPV Cycle', 'A facility that receives the monthly EPV email automatically.'],
            ['Verified (badge)', 'A facility whose Company Admin has completed the onboarding wizard.'],
        ])


def levy_summary(doc, *, technical=True):
    h1(doc, 'How the levy is calculated')
    section_intro(doc, 'The maths behind every invoice. Trade sales are what count. All amounts exclude VAT.')
    callout(doc, 'red',
            'ALL Rand amounts anywhere in EPVS are stored and displayed EXCLUSIVE of VAT. '
            'The statutory levy is calculated on the pre-VAT sales-to-trade quantities. '
            'VAT is added separately by ALS on the invoice they send you.')
    if technical:
        reference_table(doc,
            headers=['Product', 'Basis', 'Rate (excl VAT)', 'Formula'],
            rows=[
                ['Eggs',   'Dozens sold to trade', 'R 0.020 / dozen', 'SoldToTrade × 0.020'],
                ['Pulp',   'kg sold to trade → dozens at 1.7', 'R 0.020 / dozen (R 0.034 / kg)', 'PulpSoldToTrade × 1.7 × 0.020'],
                ['Powder', 'kg sold to trade', 'R 0.020 / kg', 'PowderSoldToTrade × 0.020'],
            ])
        para(doc,
             'These fields are recorded so your closing stock balances, but they DO '
             'NOT contribute to the levy: sales to staff, sales through a farm stall, '
             'transfers to other producers, pulp / powder sold to other producers, '
             'market returns, machine loss, sent-to-pulp, destroyed, exported, and '
             'conversion loss.')
    else:
        para(doc,
             'The statutory levy is worked out on what you sell to the trade only. '
             'Sales to staff, farm-stall sales, and stock transferred to other '
             'producers are recorded so that your closing stock balances, but they '
             'don\'t attract a levy.')
        bullet(doc, 'Eggs sold to trade — R 0.020 per dozen (excl VAT).')
        bullet(doc, 'Pulp sold to trade — converted at 1.7 dozens per kilogram, then charged at R 0.020 per dozen (about R 0.034 per kilogram, excl VAT).')
        bullet(doc, 'Powder sold to trade — R 0.020 per kilogram (excl VAT).')


def about_epvs(doc, *, external=False):
    h1(doc, 'About EPVS')
    section_intro(doc, 'What the platform is, who uses it, and where you fit in.')
    if external:
        para(doc,
             'The Egg Production Verification System (EPVS) is the online platform ALS '
             'Food Safety Agency uses to collect the monthly Egg Production Verifications '
             '(EPVs) from registered facilities across South Africa. It replaces the '
             'previous paper-based process. Everything happens in your web browser — no '
             'software to install and no plugins required.')
        para(doc,
             'Every month, EPVS emails your facility a link to a short online form. You '
             'complete the form, submit it, and an inspector reviews the numbers. Once '
             'they approve, ALS raises the levy invoice and you settle it in the usual '
             'way. Everything is tracked so both sides can see the current state at any time.')
    else:
        para(doc,
             'EPVS manages the statutory levy owed on graded eggs, pulp, and egg '
             'powder across every registered facility. Facilities submit monthly EPVs, '
             'inspectors verify them, and ALS invoices and reconciles payment. This '
             'manual walks through everything the current role can do.')
        para(doc, 'The five moving parts:')
        bullet(doc, 'Facilities submit their monthly EPV.')
        bullet(doc, 'Inspectors approve or capture a revised set of figures.')
        bullet(doc, 'ALS invoices and records payment.')
        bullet(doc, 'Administrators and Super Admins oversee the whole cycle.')
        bullet(doc, 'The system does the arithmetic, schedules issuance, and logs everything.')


# ══════════════════════════════════════════════════════════════════════
# ROLE MANUALS
# ══════════════════════════════════════════════════════════════════════
def build_super_admin(path):
    doc = Document()
    _configure_styles(doc)
    _footer(doc, 'Super Administrator Manual')
    cover(doc,
          title='Super Administrator Manual',
          subtitle='Full-system operations guide',
          audience='ALS Food Safety Agency Super Administrators')
    toc_page(doc)

    about_epvs(doc, external=False)
    page_break(doc)

    h1(doc, 'Logging in and getting oriented')
    section_intro(doc, 'What you see the moment you log in as a Super Administrator.')
    h2(doc, 'Login')
    para(doc, 'Open the EPVS URL in your browser. Enter your email and password. If you have forgotten your password, click "Forgot Password" for a reset link.')
    callout(doc, 'note', 'If your login fails, check the "Login History" panel on User Management — you can see the reason for every rejected attempt (wrong password, deactivated account, unknown email).')
    h2(doc, 'Landing page')
    para(doc, 'Super Admins land on the Dashboard. Nine tabs are available in the top ribbon.')
    reference_table(doc,
        headers=['Tab', 'Purpose'],
        rows=[
            ['Dashboard',            'KPI ribbon, financial charts, province overview.'],
            ['Clients',              'Consolidated Master Facility Database (CRUD + audit).'],
            ['Inspectors',           'Inspector workload, KPIs, assign facilities.'],
            ['Administrators',       'Reconciliation control centre.'],
            ['ALS Collections',      '3rd-party levy collector\'s worklist (also visible to you for transparency).'],
            ['Company Overview',     'Deep-dive per facility.'],
            ['Support',              'Support tickets across the system.'],
            ['User Management',      'Add / edit / deactivate users; login history.'],
        ])
    page_break(doc)

    h1(doc, 'Consolidated Master Facility Database')
    section_intro(doc, 'The master list of every facility, with add / edit / delete, invitations, and Send EPV.')
    h2(doc, 'The badges')
    reference_table(doc,
        headers=['Badge', 'Meaning'],
        rows=[
            ['Verified',       'A Company Admin has accepted the invitation and completed the onboarding wizard.'],
            ['On EPV Cycle',   'The facility is on the monthly auto-send schedule.'],
        ])
    h2(doc, 'Common actions')
    bullet(doc, 'Add New — opens an inline row. Business Name is required; other fields can be filled progressively.')
    bullet(doc, 'Edit — inline edit; only the changed cell is written. Every change is logged in the Change Log.')
    bullet(doc, 'Del — removes the facility AND every linked EPV, invitation, and audit row. Confirm before you click.')
    bullet(doc, 'Invite — opens the invitation modal to add a Company Admin. Their acceptance triggers the "Verified" badge.')
    bullet(doc, 'Send EPV — creates the previous-month EPV and emails all four addresses on file. First send also flips EPVCycleStatus to "On EPV Cycle" so the automatic monthly sends kick in next month.')
    bullet(doc, 'Export to Excel — downloads the current search view. Include Verified status, On EPV Cycle status, and Assigned Inspector.')
    bullet(doc, 'Change Log — flips the view to the audit trail. Every field-level change is here.')
    callout(doc, 'warning', 'Delete is irreversible without a database restore. If in doubt, "deactivate" instead by clearing the primary email so no auto-sends fire, then remove later.')
    page_break(doc)

    h1(doc, 'Company Overview')
    section_intro(doc, 'Everything about one facility on a single page.')
    h2(doc, 'Sections')
    bullet(doc, 'Company Details — editable business information.')
    bullet(doc, 'Company Users — invite, list, remove.')
    bullet(doc, 'EPVs — full monthly history. Click a row to expand: progress tracker + Levy Invoice from ALS panel when applicable.')
    bullet(doc, 'Resend Email — re-emails a Pending EPV without duplicating the record.')
    bullet(doc, 'Reconciliation — manually mark reconciled or set the received amount.')
    bullet(doc, 'Change Log — audit of every action against this facility.')
    h2(doc, 'Approve / Reject override')
    para(doc, 'You can override an inspector\'s decision by toggling the Verify column back and forth. Use sparingly — it should only be needed when an inspector made a genuine mistake and cannot correct it themselves.')
    page_break(doc)

    h1(doc, 'Dashboard')
    section_intro(doc, 'Single-glance view of every facility, verification, and payment.')
    h2(doc, 'Filters')
    para(doc, 'Filter cascades: Year → Quarter → Month. Every KPI, chart, and money figure follows the filter.')
    h2(doc, 'KPI ribbon')
    reference_table(doc,
        headers=['KPI', 'Meaning'],
        rows=[
            ['Facilities',       'Total facilities in the Master DB.'],
            ['Completed EPVs',   'Submitted EPVs in the filter window.'],
            ['Verified',         'IsVerified = 1 EPVs.'],
            ['Rejections',       'Number of Inspector EPVs (implies facility was rejected).'],
            ['Inspections Done', 'Physical inspections captured (ManualInspection).'],
            ['Outstanding',      'Facilities missing an EPV for the current period.'],
        ])
    h2(doc, 'Financial module')
    para(doc, 'Total Billed, Total Paid, Outstanding, Egg Levy, Pulp Levy (with kg + dozens), Powder Levy (with kg), Collection Rate. Two charts — Billed vs Paid month over month, and Egg vs Pulp vs Powder Levy trends.')
    h2(doc, 'Province Overview')
    para(doc, 'Facility distribution pie + per-province rejection bar chart. Handy for spotting geographic clusters of issues.')
    page_break(doc)

    h1(doc, 'Administrators (reconciliation)')
    section_intro(doc, 'The reconciliation control centre for large payment batches.')
    para(doc, 'Every completed EPV is listed with its POP, reconciled state, and amount. When a bank statement lands, work down the list ticking off each payment. Batch actions are available for multi-select.')
    callout(doc, 'note', 'Amounts entered here write to the same fields ALS uses on ALS Collections. Whichever team logs the payment first wins — the audit log records who and when.')
    page_break(doc)

    h1(doc, 'Inspectors')
    section_intro(doc, 'Inspector performance and workload management.')
    para(doc, 'Assign inspectors to facilities, view each inspector\'s KPIs (approvals, facilities visited, rejections), and jump into an EPV to override an approval where needed.')
    para(doc, 'Super Admins can also be selected as inspectors for allocation purposes when coverage is short — this lets you approve directly without creating an "Inspector" role user.')
    page_break(doc)

    h1(doc, 'ALS Collections (transparency oversight)')
    section_intro(doc, 'The 3rd-party collector\'s worksurface. You see everything they do.')
    h2(doc, 'What appears here')
    para(doc, 'Only EPVs that are inspection-cleared appear (approved by inspector, OR rejected with a completed Inspector EPV). Pending inspector reviews are hidden.')
    h2(doc, 'What you can do')
    bullet(doc, 'Everything ALS can do — mark invoice sent, upload the invoice PDF, reconcile the payment (short-pays supported), comment, view the facility\'s POP.')
    bullet(doc, 'Export to Excel — nicely-named file (`ALS Collections Report - {Month YYYY} - {date}.xlsx`) honouring your current filters.')
    h2(doc, 'The transparency log')
    para(doc, 'Every ALS-side and Super-Admin-side action is logged and shown at the bottom of every expanded row. Green badge = ALS. Indigo badge = Super Admin. Nothing is hidden from either party.')
    callout(doc, 'tip', 'Use this log during monthly reconciliation reviews with ALS. Point at the badges to show who did what, and use the "From → To" columns for audit sign-off.')
    page_break(doc)

    h1(doc, 'User Management')
    section_intro(doc, 'Adding users, changing roles, resetting passwords, and reviewing every login.')
    h2(doc, 'Adding a user')
    numbered(doc, 1, 'Click "+ Add User" on the User Management page.')
    numbered(doc, 2, 'Pick a role: Super Admin, Admin, ALS, Inspector, Company Admin, or User.')
    numbered(doc, 3, 'For Company Admin and User, pick the facility to link to.')
    numbered(doc, 4, 'Enter first name, last name, email, and set a temporary password (they can reset later).')
    numbered(doc, 5, 'Save. The user can now log in.')
    h2(doc, 'Reset password / deactivate')
    para(doc, 'Use "Reset PW" for a password change, "Deactivate" to keep the record but block login. Deactivated users can be reactivated later.')
    h2(doc, 'Login History')
    para(doc, 'The Login History button opens the audit log of every login attempt — successful and failed — with timestamp, email, IP, and reason for any failure. Use it when investigating suspicious account activity.')
    page_break(doc)

    h1(doc, 'Common workflows')
    section_intro(doc, 'Step-by-step for the tasks you do most.')
    h2(doc, 'Onboarding a new facility')
    numbered(doc, 1, 'Master Facility Database → + Add New.')
    numbered(doc, 2, 'Fill Business Name (required) + as much detail as you have.')
    numbered(doc, 3, 'Save the row.')
    numbered(doc, 4, 'Click Invite → enter the Company Admin\'s email.')
    numbered(doc, 5, 'They accept the invite → complete the wizard → the "Verified" badge appears.')
    numbered(doc, 6, 'Click Send EPV to issue their first EPV and put them on the monthly cycle.')

    h2(doc, 'Overriding an inspector\'s decision')
    numbered(doc, 1, 'Company Overview → search for the facility → open the EPV row.')
    numbered(doc, 2, 'Toggle the Verify checkbox. The change is logged.')
    numbered(doc, 3, 'Add a comment via Inspector Comment so the reason is on the record.')

    h2(doc, 'Monthly reconciliation batch')
    numbered(doc, 1, 'Administrators tab → sort by POP Uploaded date.')
    numbered(doc, 2, 'Cross-reference the bank statement.')
    numbered(doc, 3, 'For each matched EPV, tick Reconciled and enter the received amount.')
    numbered(doc, 4, 'Short-pays: enter the amount received. Outstanding remains visible on the Dashboard.')
    page_break(doc)

    h1(doc, 'Troubleshooting')
    section_intro(doc, 'What to try when something looks wrong.')
    reference_table(doc,
        headers=['Symptom', 'Likely cause', 'What to try'],
        rows=[
            ['A facility isn\'t receiving EPV emails.', 'The four facility addresses are wrong or bouncing.', 'Company Overview → Company Details → confirm every email. Check EmailSendLog for delivery status.'],
            ['A facility says they can\'t submit for last month.', 'The current-month lockout blocks the CURRENT month, not last month.', 'Confirm the period the facility is trying to submit. If it is genuinely last month, they should be able to submit.'],
            ['Verified badge missing though they\'ve accepted the invite.', 'They created the account but didn\'t finish the wizard.', 'Ask them to log back in and complete the onboarding steps. Or check Invitations directly.'],
            ['Duplicate EPV attempts.', 'Someone clicked Send EPV twice in quick succession.', 'The system rejects the duplicate. Check the EPV list — only one row should exist per period.'],
        ])
    page_break(doc)

    support_and_glossary(doc, external=False)
    page_break(doc)
    levy_summary(doc, technical=True)
    doc.save(path)


def build_admin(path):
    doc = Document()
    _configure_styles(doc)
    _footer(doc, 'Administrator Manual')
    cover(doc,
          title='Administrator Manual',
          subtitle='Operational oversight and reconciliation',
          audience='ALS Food Safety Agency Administrators')
    toc_page(doc)

    about_epvs(doc, external=False)
    page_break(doc)

    h1(doc, 'Logging in and getting oriented')
    section_intro(doc, 'What you see the moment you log in as an Administrator.')
    para(doc, 'Administrators land on the Dashboard. Seven tabs are available.')
    reference_table(doc,
        headers=['Tab', 'Purpose'],
        rows=[
            ['Dashboard',        'KPI ribbon, financial charts, province overview.'],
            ['Clients',          'Consolidated Master Facility Database (CRUD + audit).'],
            ['Inspectors',       'Inspector workload and allocation.'],
            ['Administrators',   'Reconciliation control centre — where you spend most of your day.'],
            ['Company Overview', 'Deep-dive per facility.'],
            ['Support',          'Support tickets across the system.'],
            ['User Management',  'Add / edit / deactivate users; login history.'],
        ])
    callout(doc, 'note', 'ALS Collections is restricted to ALS staff and Super Administrators. You will not see that tab.')
    page_break(doc)

    h1(doc, 'Consolidated Master Facility Database')
    section_intro(doc, 'Add and manage facilities.')
    para(doc, 'You have the same abilities as Super Administrators here — add, edit, delete facilities; send EPVs; export.')
    h2(doc, 'Badges')
    reference_table(doc,
        headers=['Badge', 'Meaning'],
        rows=[
            ['Verified',     'Company Admin has completed the onboarding wizard.'],
            ['On EPV Cycle', 'Facility is on the monthly auto-send schedule.'],
        ])
    h2(doc, 'Export to Excel')
    para(doc, 'The Export to Excel button downloads the current search view, including contact details, verified status, cycle status, and assigned inspector. Filename: `Consolidated Master Facility Database - YYYY-MM-DD.xlsx`.')
    page_break(doc)

    h1(doc, 'Company Overview')
    section_intro(doc, 'One facility at a time.')
    para(doc, 'From the EPV list you can Resend the Pending email, mark manual reconciliation, and view any invoice ALS has uploaded to the facility.')
    para(doc, 'The Change Log shows every action against this facility for full accountability.')
    page_break(doc)

    h1(doc, 'Dashboard')
    section_intro(doc, 'Your single-glance view.')
    para(doc, 'Filterable by Year / Quarter / Month. Every widget follows the filter.')
    h2(doc, 'The KPIs to watch')
    reference_table(doc,
        headers=['KPI', 'Why it matters'],
        rows=[
            ['Completed EPVs',   'Volume of monthly submissions received.'],
            ['Verified',         'How many are ready for invoicing by ALS.'],
            ['Rejections',       'Where the inspection process caught issues.'],
            ['Outstanding',      'Facilities that missed a submission this period.'],
            ['Collection Rate',  'How much of what was billed has actually been paid.'],
        ])
    page_break(doc)

    h1(doc, 'Administrators (reconciliation)')
    section_intro(doc, 'Your daily worksurface.')
    para(doc, 'Every completed EPV appears here with its POP, reconciled status, and amount. Use it to reconcile in batches when payment files arrive.')
    h2(doc, 'A typical reconciliation session')
    numbered(doc, 1, 'Download the bank statement CSV.')
    numbered(doc, 2, 'On Administrators, sort by POP Uploaded date descending.')
    numbered(doc, 3, 'Match each bank line to an EPV by reference number or amount.')
    numbered(doc, 4, 'Tick Reconciled and enter the amount received.')
    numbered(doc, 5, 'Short pays — enter what was actually received. Outstanding stays visible on the Dashboard.')
    callout(doc, 'tip', 'Use Company Overview to view the POP inline before reconciling. This helps confirm the reference on the bank statement matches the EPV.')
    page_break(doc)

    h1(doc, 'Inspectors')
    section_intro(doc, 'Assign, monitor, and support the inspection team.')
    para(doc, 'View inspector workload, KPIs, and progress. Assign inspectors to facilities so incoming EPVs land in the right queue.')
    para(doc, 'Super Admins can also be selected as inspectors for coverage — treat that as a stop-gap rather than a permanent assignment.')
    page_break(doc)

    h1(doc, 'User Management')
    section_intro(doc, 'Managing internal users.')
    para(doc, 'Add and manage internal users — Inspectors, other Administrators. Reset passwords and deactivate accounts. The Login History button shows every authentication attempt for audit.')
    page_break(doc)

    h1(doc, 'Common workflows')
    section_intro(doc, 'The tasks you do most.')
    h2(doc, 'Onboarding a facility')
    numbered(doc, 1, 'Master Facility Database → + Add New.')
    numbered(doc, 2, 'Fill Business Name + as many contacts as you have.')
    numbered(doc, 3, 'Invite the Company Admin.')
    numbered(doc, 4, 'When they finish onboarding, click Send EPV.')
    h2(doc, 'End-of-day reconciliation')
    para(doc, 'Filter Administrators for today\'s payments. Confirm each is reconciled. Escalate mismatches via a Support ticket.')
    page_break(doc)

    support_and_glossary(doc, external=False)
    page_break(doc)
    levy_summary(doc, technical=True)
    doc.save(path)


def build_inspector(path):
    doc = Document()
    _configure_styles(doc)
    _footer(doc, 'Inspector Manual')
    cover(doc,
          title='Inspector Manual',
          subtitle='Verifying EPVs and managing your facilities',
          audience='ALS Food Safety Agency Inspectors')
    toc_page(doc)

    about_epvs(doc, external=False)
    page_break(doc)

    h1(doc, 'Logging in and getting oriented')
    section_intro(doc, 'What you see the moment you log in as an Inspector.')
    para(doc, 'You land on the Inspectors page. Four tabs are available.')
    reference_table(doc,
        headers=['Tab', 'Purpose'],
        rows=[
            ['Inspectors',       'Your dashboard: KPIs, pending approvals, upcoming visits.'],
            ['Company Overview', 'Deep-dive per facility, including facilities you cover.'],
            ['Support',          'Raise or track support tickets.'],
            ['User Management',  'Your own profile — change password, review your login history.'],
        ])
    page_break(doc)

    h1(doc, 'The Inspectors dashboard')
    section_intro(doc, 'Your daily worksurface.')
    h2(doc, 'KPIs')
    reference_table(doc,
        headers=['KPI', 'Meaning'],
        rows=[
            ['Approvals Actioned', 'Completed EPVs you have approved or rejected.'],
            ['Facilities Visited', 'Manual inspections logged in the current quarter.'],
            ['Rejections',         'Facility EPVs you rejected and captured a revised set for.'],
        ])
    h2(doc, 'Pending Approvals')
    para(doc, 'The primary list on this page. Any completed facility EPV that has not yet been approved or rejected shows here. Click through to review the numbers.')
    h2(doc, 'Assigned facilities')
    para(doc, 'The facilities allocated to you appear in your list. If you need to cover another inspector\'s workload, ask an Administrator to update the allocation.')
    page_break(doc)

    h1(doc, 'Verifying an EPV (approve)')
    section_intro(doc, 'The happy path — the facility got the numbers right.')
    numbered(doc, 1, 'Open the EPV from your dashboard or from Company Overview.')
    numbered(doc, 2, 'Walk through the four wizard steps: Business Details, Eggs, Pulp, Powder.')
    numbered(doc, 3, 'Check the Review step for the theoretical vs actual closing stock. A variance reason should be attached if they differ.')
    numbered(doc, 4, 'Check the purchase evidence attachments — supplier invoices, delivery notes.')
    numbered(doc, 5, 'On the EPV row, tick Approve. The status becomes "Approved" and the EPV is released to ALS for invoicing.')
    callout(doc, 'tip', 'Add a comment via Inspector Comment before approving. Anyone reading the record months later will thank you.')
    page_break(doc)

    h1(doc, 'Rejecting an EPV and capturing revised figures')
    section_intro(doc, 'When the facility\'s numbers don\'t match what you observed.')
    numbered(doc, 1, 'On the EPV row, click Reject on the Verify column.')
    numbered(doc, 2, 'A new Inspector EPV opens, pre-filled with the facility\'s figures.')
    numbered(doc, 3, 'Change what needs changing. Wizard steps mirror the facility EPV.')
    numbered(doc, 4, 'Attach any supporting evidence.')
    numbered(doc, 5, 'Add a comment explaining what changed and why. This is the most important part of a rejection.')
    numbered(doc, 6, 'Submit. ALS will invoice on YOUR figures, not the facility\'s.')
    callout(doc, 'warning', 'Once submitted, the Inspector EPV is locked. A Super Admin can override if a fix is needed — raise a support ticket describing what needs to change.')
    page_break(doc)

    h1(doc, 'Company Overview')
    section_intro(doc, 'Deep-dive on any facility you cover.')
    para(doc, 'Full monthly EPV history, POP uploads, reconciliation status, and any invoice ALS has uploaded to the facility. Use this to spot patterns — repeat rejections, chronic short-payers, month-over-month trends.')
    page_break(doc)

    h1(doc, 'Common workflows')
    section_intro(doc, 'The tasks you do most.')
    h2(doc, 'Reviewing your monthly workload')
    numbered(doc, 1, 'Inspectors dashboard → sort Pending Approvals by SentAt.')
    numbered(doc, 2, 'Work oldest-first so nothing ages more than a few days.')
    numbered(doc, 3, 'For anything you can\'t clear same-day, add a comment to the EPV so ALS and admins know why.')
    h2(doc, 'On-site physical inspection')
    numbered(doc, 1, 'Company Overview → open the facility.')
    numbered(doc, 2, 'On the EPV row for the period visited, tick Manual Inspection.')
    numbered(doc, 3, 'Add any observations to the Inspector Comment.')
    page_break(doc)

    h1(doc, 'Troubleshooting')
    section_intro(doc, 'When something looks off.')
    reference_table(doc,
        headers=['Symptom', 'Likely cause', 'What to try'],
        rows=[
            ['A facility\'s figures look impossible.', 'Wrong opening stock brought forward.', 'Compare the previous month\'s ClosingStock with this month\'s OpeningStock. Raise a support ticket if they don\'t match.'],
            ['An EPV isn\'t appearing in Pending Approvals.', 'Not yet submitted, or already approved.', 'Company Overview → find the EPV. Check its Status and IsVerified.'],
            ['A rejected EPV needs another change.', 'Inspector EPVs lock on submit.', 'Raise a support ticket for a Super Admin override.'],
        ])
    page_break(doc)

    support_and_glossary(doc, external=False)
    page_break(doc)
    levy_summary(doc, technical=True)
    doc.save(path)


def build_als(path):
    doc = Document()
    _configure_styles(doc)
    _footer(doc, 'ALS Collections Manual')
    cover(doc,
          title='ALS Collections Manual',
          subtitle='Levy invoicing, reconciliation and reporting',
          audience='ALS Collections team')
    toc_page(doc)

    about_epvs(doc, external=False)
    page_break(doc)

    h1(doc, 'Logging in and getting oriented')
    section_intro(doc, 'What you see the moment you log in.')
    para(doc, 'You land directly on the ALS Collections page. It is your primary worksurface. The top ribbon also shows Company Overview (read-only across facilities), Support, and your own Settings.')
    callout(doc, 'note', 'Only ALS users and Super Administrators can access the ALS Collections page. Every other role is redirected elsewhere on login.')
    page_break(doc)

    h1(doc, 'The ALS Collections page')
    section_intro(doc, 'What appears here, and why nothing else does.')
    h2(doc, 'What appears in the worklist')
    para(doc, 'Only EPVs that have been inspection-cleared appear. An EPV shows up in one of two states:')
    reference_table(doc,
        headers=['State', 'What it means', 'Invoice amount'],
        rows=[
            ['Approved',    'Inspector approved the facility\'s submitted figures.', 'Facility\'s Egg + Pulp + Powder levies.'],
            ['Rejected → Insp', 'Inspector rejected the facility\'s figures and captured a revised set.', 'Inspector\'s Egg + Pulp + Powder levies.'],
        ])
    callout(doc, 'tip', 'EPVs still waiting on the inspector are deliberately hidden — you only see EPVs you can act on today.')
    h2(doc, 'The KPI ribbon')
    reference_table(doc,
        headers=['KPI', 'Meaning'],
        rows=[
            ['Invoiceable',      'Total EPVs waiting to be invoiced.'],
            ['Invoices Sent',    'How many have been marked "sent".'],
            ['Reconciled',       'How many have been marked as paid.'],
            ['Total Billable',   'What you should collect for the filter.'],
            ['Total Collected',  'What has been received (matches ReconciledAmount).'],
            ['Outstanding',      'Billable − Collected.'],
            ['Collection Rate',  'Collected / Billable as a percentage.'],
        ])
    page_break(doc)

    h1(doc, 'Filtering the worklist')
    section_intro(doc, 'Slicing the list to focus on what matters right now.')
    para(doc, 'The filter bar above the table gives you six controls:')
    bullet(doc, 'Free-text search — facility name, account code, or reference number.')
    bullet(doc, 'Year and Month — restrict to a specific period.')
    bullet(doc, 'Approvals — Approved only, Rejected only, or All.')
    bullet(doc, 'Invoices — Sent, Not sent, or All.')
    bullet(doc, 'Payments — Reconciled, Unreconciled, or All.')
    para(doc, 'The KPI ribbon and the Excel export both honour whatever filters you have on.')
    page_break(doc)

    h1(doc, 'The expanded row')
    section_intro(doc, 'Click any row to reveal the full worksurface for that EPV.')
    para(doc, 'The expanded panel has four columns:')
    h2(doc, 'Amount breakdown')
    para(doc, 'Facility figures, Inspector figures (if any), and the resolved invoice amount. If the inspector rejected the facility, this is where you see both sets of numbers side by side.')
    h2(doc, 'ALS invoice PDF')
    para(doc, 'Upload, view, or remove your invoice. PDF / PNG / JPG, up to 15 MB. Uploading a new file replaces the previous one. Removed files are deleted from disk immediately.')
    h2(doc, 'Payment')
    para(doc, 'Amount owed at the top. Below: link to the facility POP if they uploaded one, then the reconcile controls (amount received + Mark paid button).')
    h2(doc, 'ALS comment')
    para(doc, 'Free text, visible to ALS and Super Admins. Use it as your working log for the EPV — chase reminders, call notes, escalation status.')
    page_break(doc)

    h1(doc, 'Marking an invoice as sent')
    section_intro(doc, 'When you\'ve dispatched the invoice from your billing system.')
    para(doc, 'Tick the "Invoice Sent" checkbox on the row. The system records the timestamp and who marked it. The facility sees a "Levy Invoice from ALS" panel appear inside their Company Overview the moment you upload the PDF (below).')
    page_break(doc)

    h1(doc, 'Uploading your invoice PDF')
    section_intro(doc, 'Making the invoice visible to the facility.')
    numbered(doc, 1, 'Expand the EPV row.')
    numbered(doc, 2, 'Click "Upload invoice" in the ALS invoice PDF column.')
    numbered(doc, 3, 'Pick a PDF, PNG, or JPG up to 15 MB.')
    numbered(doc, 4, 'The uploaded file is visible immediately in the row (View link) and inside the facility\'s Company Overview.')
    callout(doc, 'tip', 'Use consistent file naming — e.g. "ALS-INV-{ref}.pdf". Makes reconciliation later much easier for both sides.')
    page_break(doc)

    h1(doc, 'Reconciling a payment')
    section_intro(doc, 'When money lands and needs to be matched to an EPV.')
    para(doc, 'Reconciliation writes to the shared reconciled fields the whole system already uses — nothing lives in an ALS-only silo. Once you reconcile, the Dashboard, Administrators tab, and Company Overview all reflect it automatically.')
    h2(doc, 'A full payment')
    numbered(doc, 1, 'Expand the row.')
    numbered(doc, 2, 'The amount received input pre-fills with the invoice amount.')
    numbered(doc, 3, 'Click "Mark paid".')
    numbered(doc, 4, 'Outstanding drops to R 0.00. The row highlights green.')
    h2(doc, 'A short pay')
    numbered(doc, 1, 'Expand the row.')
    numbered(doc, 2, 'Type the actual amount received (e.g. 800 for an R 1000 invoice).')
    numbered(doc, 3, 'Click "Mark paid" — the row is reconciled but Outstanding shows R 200.')
    numbered(doc, 4, 'Follow up with the facility for the balance. Add a comment on the row so the reason is on the record.')
    h2(doc, 'Reverting a reconciliation')
    para(doc, 'If you reconciled in error, click "Clear" to reset. Reconciled amount and the audit stamps clear together.')
    callout(doc, 'tip', 'If the facility uploaded their POP, click "View POP" first to confirm the amount before reconciling.')
    page_break(doc)

    h1(doc, 'The transparency history')
    section_intro(doc, 'A shared audit trail with Super Admins.')
    para(doc, 'Every action in this page is logged for both Super Admins and ALS to see. The bottom of every expanded row shows a "History · transparency log" table.')
    reference_table(doc,
        headers=['Column', 'Meaning'],
        rows=[
            ['When',   'Timestamp of the action.'],
            ['Action', 'Invoice Sent, Invoice Uploaded, Reconciled Payment, or Comment.'],
            ['From',   'Value before the action.'],
            ['To',     'Value after the action.'],
            ['By',     'Who did it, with a role badge — green for ALS, indigo for Super Admin.'],
        ])
    callout(doc, 'note', 'The log is shared. Super Admins can see what ALS has done and vice versa. Nothing is hidden between the two sides — that\'s intentional.')
    page_break(doc)

    h1(doc, 'Exporting to Excel')
    section_intro(doc, 'Getting the worklist out for offline analysis or hand-off.')
    para(doc, 'The Export to Excel button downloads the current filtered view. Filename pattern:')
    para(doc, '    ALS Collections Report - {Month YYYY if filtered} - YYYY-MM-DD.xlsx')
    para(doc, 'Every column visible in the table is included, plus:')
    bullet(doc, 'Facility vs Inspector breakdown per levy (egg, pulp, powder).')
    bullet(doc, 'POP uploaded flag.')
    bullet(doc, 'Amount paid and outstanding.')
    bullet(doc, 'Your ALS comment for the row.')
    page_break(doc)

    h1(doc, 'Company Overview')
    section_intro(doc, 'Read-only cross-facility.')
    para(doc, 'Use it when a facility asks you a question — you can quickly see their EPV history, payment status, and any comments the inspection team has left. You cannot edit facility data or capture EPVs from here.')
    page_break(doc)

    h1(doc, 'Common workflows')
    section_intro(doc, 'The tasks you do most.')
    h2(doc, 'Sending the monthly invoices')
    numbered(doc, 1, 'ALS Collections → filter to this month\'s invoiceable set.')
    numbered(doc, 2, 'Generate the invoice from your billing system for each row.')
    numbered(doc, 3, 'Upload the invoice PDF on the row.')
    numbered(doc, 4, 'Tick "Invoice Sent" — the facility now sees it in their Company Overview.')
    h2(doc, 'Weekly reconciliation')
    numbered(doc, 1, 'Filter to Reconciled = No, Invoice Sent = Yes.')
    numbered(doc, 2, 'Cross-reference the bank feed / receipts against outstanding invoices.')
    numbered(doc, 3, 'Enter each amount received and tick "Mark paid".')
    numbered(doc, 4, 'Add a comment for anything unusual (short pay, part payment, EFT reference mismatch).')
    h2(doc, 'Monthly close')
    numbered(doc, 1, 'Filter to the closing month.')
    numbered(doc, 2, 'Export to Excel. Save it for records.')
    numbered(doc, 3, 'Review the KPI ribbon — Collection Rate for the period should be greater than 90%.')
    page_break(doc)

    h1(doc, 'Troubleshooting')
    section_intro(doc, 'When something looks off.')
    reference_table(doc,
        headers=['Symptom', 'Likely cause', 'What to try'],
        rows=[
            ['A completed EPV isn\'t appearing here.', 'The inspector hasn\'t approved or rejected it yet.', 'Ask the inspection team to review. Only inspection-cleared EPVs appear.'],
            ['Invoice amount seems too low.', 'The inspector rejected and reduced the figures.', 'Expand the row — "Rejected → Insp" appears in the Approval column. The Inspector figures are shown alongside the facility\'s.'],
            ['Facility says they never got the invoice.', 'Their four email addresses may be out of date.', 'Company Overview → confirm addresses. If correct, ask an Administrator to check EmailSendLog.'],
            ['Upload fails with "file too large".', 'The invoice PDF is over 15 MB.', 'Compress the PDF (any online tool works). 15 MB is the platform-wide cap.'],
        ])
    page_break(doc)

    support_and_glossary(doc, external=False)
    page_break(doc)
    levy_summary(doc, technical=True)
    doc.save(path)


def build_company_admin(path):
    doc = Document()
    _configure_styles(doc)
    _footer(doc, 'Company Administrator Guide')
    cover(doc,
          title='Company Administrator Guide',
          subtitle='Managing your facility on the EPVS platform',
          audience='Facility Company Administrators')
    toc_page(doc)

    h1(doc, 'Welcome')
    section_intro(doc, 'What this guide covers and how to use it.')
    para(doc, 'Welcome to the Egg Production Verification System (EPVS). This guide walks you through everything you need to do inside the platform: accepting your invitation, submitting your monthly EPV, uploading your Proof of Payment, viewing invoices, and adding users from your team.')
    para(doc, 'Everything happens in your web browser — there is no software to install. Use it on a laptop or desktop. The mobile view works, but the EPV form is easiest on a full-size screen.')
    page_break(doc)

    about_epvs(doc, external=True)
    page_break(doc)

    h1(doc, 'Getting started — accepting your invitation')
    section_intro(doc, 'From the email you received to your first login.')
    para(doc, 'Your invitation arrives by email. The link inside is personal to you and expires after 30 days.')
    numbered(doc, 1, 'Open the invitation email and click "Accept Invitation".')
    numbered(doc, 2, 'Enter your first name, last name, and choose a password.')
    numbered(doc, 3, 'You will be guided through a short onboarding wizard to confirm your facility details — business info, owner, accounts contact, facility manager, physical address, VAT number.')
    numbered(doc, 4, 'When you finish the wizard, your facility is marked "Verified". You land on your Company Overview page.')
    callout(doc, 'note', 'If your link has expired, ask ALS to resend it. You can raise a support ticket for that from the login page.')
    h2(doc, 'Setting a strong password')
    para(doc, 'Choose a password of at least eight characters with a mix of letters, numbers, and symbols. If you forget it later, "Forgot Password" on the login page sends a reset link.')
    page_break(doc)

    h1(doc, 'Your Company Overview')
    section_intro(doc, 'Your home page after login.')
    para(doc, 'The Company Overview shows your facility details at the top and every EPV underneath, from your very first submission to the most recent. It is where you spend most of your time.')
    h2(doc, 'What you can do here')
    reference_table(doc,
        headers=['Section', 'What you can do'],
        rows=[
            ['Company Details',   'Update your facility information when it changes — contact numbers, addresses, VAT number.'],
            ['Company Users',     'Invite and manage users on your team.'],
            ['EPVs',              'Open any past EPV to view or continue where you left off.'],
            ['POP (per EPV row)', 'Upload a Proof of Payment once you have paid ALS.'],
            ['Levy Invoice from ALS (per EPV row)', 'Download the invoice ALS has sent you.'],
            ['Change Log',        'See a full history of changes to your facility.'],
        ])
    callout(doc, 'tip', 'Keep your contact details up to date. EPV emails and payment reminders go to the four addresses on file: primary, owner, accounts, and facility manager. If any of those bounce, you might miss a notification.')
    page_break(doc)

    h1(doc, 'Completing your monthly EPV')
    section_intro(doc, 'The 4-step wizard, in detail.')
    para(doc, 'At the start of every month, you receive an email asking you to complete the EPV for the previous month. Click the button in the email to open the form.')
    h2(doc, 'Step 1 — Business Details')
    para(doc, 'The form pre-fills your business details from your facility record. Confirm the trading name, name of the authorised person completing the form, their position, phone number, cell number, and email. Correct anything that has changed.')
    h2(doc, 'Step 2 — Eggs')
    para(doc, 'Enter values in dozens.')
    bullet(doc, 'Opening Stock — how many dozens you had at the start of the month (carried over from last month\'s closing stock).')
    bullet(doc, 'Eggs Produced During the Month — production for the period.')
    bullet(doc, 'Purchases — Graded, Ungraded, and Market Returns.')
    bullet(doc, 'Deductions — Machine Loss, Sent to Pulp, Destroyed, Eggs Exported.')
    bullet(doc, 'Sales — Sold to Trade (this is what drives your levy), Sold to Staff / Own Use, Sold through Farm Stall.')
    bullet(doc, 'Transferred — to other producers.')
    bullet(doc, 'Actual Closing Stock — what you counted at month-end.')
    h2(doc, 'Step 3 — Pulp')
    para(doc, 'Values in kilograms. Same structure as eggs. Sold to Trade drives the pulp levy (converted at 1.7 dozens per kilogram, then R 0.020 per dozen).')
    h2(doc, 'Step 4 — Powder')
    para(doc, 'Values in kilograms. Same structure as pulp. Sold to Trade drives the powder levy at R 0.020 per kilogram.')
    h2(doc, 'The Review step')
    para(doc, 'Every total is calculated for you. Check that your theoretical closing stock roughly matches your actual — if it doesn\'t, a Variance Reason box will appear. Fill it in briefly (e.g. "physical count differed by 200 dozen — breakage during transit").')
    callout(doc, 'warning', 'You can only submit an EPV for a month that has already ended. The current month is not available until the month closes. This is deliberate — the platform needs a full month to work with.')
    page_break(doc)

    h1(doc, 'Purchase evidence — comments and attachments')
    section_intro(doc, 'What to include when you purchased eggs, pulp, or powder.')
    para(doc, 'When you enter a value for Graded Eggs Purchased, Ungraded Eggs Purchased, Pulp Purchased, or Powder Purchased, a "Source & Supplier Detail" box appears in the wizard. Please:')
    bullet(doc, 'Describe where and from whom you purchased.')
    bullet(doc, 'Attach the supplier invoice or delivery note.')
    para(doc, 'Accepted formats: PDF, PNG, JPG. Maximum file size: 15 MB per attachment.')
    callout(doc, 'note', 'The attachments help your inspector complete their review faster and reduce the chance of the EPV being sent back for clarification.')
    page_break(doc)

    h1(doc, 'After you submit')
    section_intro(doc, 'What happens next, and what you need to do.')
    para(doc, 'Once you click Submit, your EPV is locked. An inspector reviews it — usually within a few working days. Two possible outcomes:')
    bullet(doc, 'If your numbers are approved, ALS raises the invoice and you receive it.')
    bullet(doc, 'If the inspector adjusts the numbers, ALS will invoice on the adjusted values. You will see the reason in the Inspector Comment.')
    para(doc, 'Either way, when the invoice is ready you can download it from your Company Overview.')
    h2(doc, 'Uploading your Proof of Payment')
    numbered(doc, 1, 'On the EPV row, click Upload POP.')
    numbered(doc, 2, 'Choose your bank confirmation PDF (or PNG / JPG). Up to 15 MB.')
    numbered(doc, 3, 'Add a short comment if useful — e.g. bank reference, "part of batch March".')
    numbered(doc, 4, 'ALS marks the EPV as reconciled once they have confirmed the payment.')
    h2(doc, 'Downloading your invoice from ALS')
    para(doc, 'When ALS uploads your invoice, a green "Levy Invoice from ALS" panel appears in the expanded view of the EPV row. Click View to download.')
    page_break(doc)

    h1(doc, 'Managing users from your team')
    section_intro(doc, 'Adding colleagues who need access to your facility.')
    para(doc, 'On the Users section of your Company Overview you can invite additional users. Two roles are available for company staff:')
    reference_table(doc,
        headers=['Role', 'Can do'],
        rows=[
            ['Company Admin', 'Manage the facility, invite other users, submit EPVs, upload POPs.'],
            ['User',          'Submit EPVs and upload POPs. Cannot invite others.'],
        ])
    para(doc, 'Invited users receive an email link and set their own password.')
    callout(doc, 'tip', 'Have at least two Company Admins on record for continuity. If your primary contact is on leave, the second can still act.')
    page_break(doc)

    h1(doc, 'Your account')
    section_intro(doc, 'Password, profile, and this manual.')
    para(doc, 'Click your name in the top-right corner and select User Management to update your own details or change your password. If you forget your password, click "Forgot Password" on the login page for a reset link.')
    para(doc, 'The User Manual button in the top ribbon downloads this guide as a PDF at any time for offline reading.')
    page_break(doc)

    h1(doc, 'Common workflows')
    section_intro(doc, 'Answers to the "how do I…" questions you\'ll get from your team.')
    h2(doc, 'A new team member needs access.')
    numbered(doc, 1, 'Company Overview → Users → Invite.')
    numbered(doc, 2, 'Pick their role (Company Admin or User).')
    numbered(doc, 3, 'Enter their email → send.')
    h2(doc, 'A previous EPV had an error.')
    numbered(doc, 1, 'Raise a support ticket from any page in the platform.')
    numbered(doc, 2, 'Describe the change needed and reference the EPV number.')
    numbered(doc, 3, 'An administrator will unlock and correct the EPV, or issue a corrective entry.')
    h2(doc, 'You need to change your facility name / address / VAT.')
    numbered(doc, 1, 'Company Overview → Company Details → Edit.')
    numbered(doc, 2, 'Change what needs to change.')
    numbered(doc, 3, 'Save. Every change is logged in the Change Log.')
    page_break(doc)

    h1(doc, 'Troubleshooting')
    section_intro(doc, 'Quick answers to common issues.')
    reference_table(doc,
        headers=['Question', 'Answer'],
        rows=[
            ['I did not get the EPV email this month.',           'Check spam. If nothing there, log in — the EPV is still available on your Company Overview even if the email did not arrive. Also, ask ALS to resend from Company Overview → Resend Email.'],
            ['The wizard won\'t let me submit for this month.',   'You can only submit for months that have already ended. The current month is not available until it closes.'],
            ['I made a typo and already submitted.',              'Raise a support ticket. Administrators can unlock and correct.'],
            ['My colleague can\'t see my facility.',              'They need an invitation from you. Company Overview → Users → Invite.'],
            ['Where is the ALS invoice?',                         'Expand the EPV row on your Company Overview. The Levy Invoice from ALS panel appears once ALS uploads it.'],
            ['The wizard is not loading anything.',               'Check that your browser is up to date (Chrome, Edge, Firefox latest). Clear cache and try again.'],
        ])
    page_break(doc)

    support_and_glossary(doc, external=True)
    page_break(doc)
    levy_summary(doc, technical=False)
    doc.save(path)


def build_user(path):
    doc = Document()
    _configure_styles(doc)
    _footer(doc, 'User Guide')
    cover(doc,
          title='User Guide',
          subtitle='Submitting EPVs and viewing your facility',
          audience='Facility Users')
    toc_page(doc)

    h1(doc, 'Welcome')
    section_intro(doc, 'What this guide covers, in plain language.')
    para(doc, 'Welcome to the Egg Production Verification System (EPVS). This guide walks you through the monthly submissions and payments for your facility. It is short on purpose — if you can fill in a form online, you can use EPVS.')
    page_break(doc)

    about_epvs(doc, external=True)
    page_break(doc)

    h1(doc, 'Getting started')
    section_intro(doc, 'From the invitation email to your first login.')
    para(doc, 'You receive an invitation email from your Company Administrator. The link is personal to you and expires after 30 days.')
    numbered(doc, 1, 'Open the email and click "Accept Invitation".')
    numbered(doc, 2, 'Fill in your first name, last name, and choose a password.')
    numbered(doc, 3, 'You land on the Company Overview page.')
    callout(doc, 'note', 'If your link has expired, ask your Company Admin to resend it.')
    page_break(doc)

    h1(doc, 'Company Overview — your home page')
    section_intro(doc, 'The page you land on after logging in.')
    para(doc, 'Every EPV for your facility, from oldest to newest, with the current status of each one. You\'ll receive an email at the start of every month asking you to complete the EPV for the previous month. Click the button in the email and follow the wizard.')
    page_break(doc)

    h1(doc, 'Completing your monthly EPV')
    section_intro(doc, 'The 4-step wizard, in plain language.')
    para(doc, 'The form is a short wizard.')
    h2(doc, 'Step 1 — Business details')
    para(doc, 'Confirm the details for the person completing the form.')
    h2(doc, 'Step 2 — Eggs')
    para(doc, 'Enter opening stock, purchases, deductions, sales, and actual closing stock — all in dozens.')
    h2(doc, 'Step 3 — Pulp')
    para(doc, 'Same idea as eggs, but the values are in kilograms.')
    h2(doc, 'Step 4 — Powder')
    para(doc, 'Same idea again, in kilograms.')
    h2(doc, 'The Review step')
    para(doc, 'The system does the arithmetic for you and shows the totals. Confirm and submit. If the theoretical closing stock differs from what you actually counted, a Variance Reason box appears — briefly explain why (e.g. breakage during transit).')
    h2(doc, 'Purchase evidence')
    para(doc, 'When you enter a value for eggs, pulp, or powder purchased, a source & supplier box appears. Describe where and from whom you purchased, and attach the supplier invoice or delivery note (PDF, PNG, or JPG up to 15 MB).')
    callout(doc, 'warning', 'You can only submit an EPV for a month that has ended. The current month is not available until it closes.')
    page_break(doc)

    h1(doc, 'After you submit')
    section_intro(doc, 'What happens next.')
    para(doc, 'An inspector reviews your EPV. Once approved, ALS sends the invoice. Pay it as you normally would, then upload the Proof of Payment on the EPV row on your Company Overview. Add a short comment if useful (e.g. bank reference).')
    para(doc, 'When ALS confirms your payment, the EPV is marked as reconciled. You will also see the invoice PDF that ALS uploaded in the expanded view.')
    callout(doc, 'tip', 'Upload the POP as soon as you have paid. It speeds up reconciliation and keeps your facility\'s payment history clean.')
    page_break(doc)

    h1(doc, 'Your account')
    section_intro(doc, 'Password, profile, and this manual.')
    para(doc, 'Click your name in the top-right corner to update your details or change your password. If you forget your password, click "Forgot Password" on the login page for a reset link.')
    para(doc, 'The User Manual button in the top ribbon downloads this guide as a PDF for offline reading.')
    page_break(doc)

    h1(doc, 'Troubleshooting')
    section_intro(doc, 'Quick answers to common issues.')
    reference_table(doc,
        headers=['Question', 'Answer'],
        rows=[
            ['I did not get the monthly email.', 'Check spam. If nothing, log in and use the EPV on your Company Overview page — the email is a reminder, not the only way in.'],
            ['I submitted the wrong figures.',   'Raise a support ticket. An administrator can unlock and correct.'],
            ['Where is my invoice from ALS?',     'Expand the EPV row on your Company Overview. The Levy Invoice from ALS panel appears once ALS uploads it.'],
        ])
    page_break(doc)

    support_and_glossary(doc, external=True)
    page_break(doc)
    levy_summary(doc, technical=False)
    doc.save(path)


def build_process_flow(path):
    doc = Document()
    _configure_styles(doc)
    _footer(doc, 'Process Flow Documentation')
    cover(doc,
          title='Process Flow Documentation',
          subtitle='End-to-end EPV lifecycle across every role',
          audience='System documentation — internal reference')
    toc_page(doc)

    h1(doc, 'Overview')
    section_intro(doc, 'How a single EPV moves from issuance to reconciliation.')
    para(doc, 'This document maps how a single EPV moves through the platform from issuance to reconciliation. Every role has a defined step and every step is logged.')
    page_break(doc)

    h1(doc, 'The five stages')

    h2(doc, 'Stage 1 — Issuance')
    para(doc, 'An EPV can be issued in two ways.')
    bullet(doc, 'Manual send from a Super Admin or Administrator on the Master Facility Database or Company Overview.')
    bullet(doc, 'Automatic monthly send by the scheduler for every facility marked "On EPV Cycle". Runs daily; only creates the EPV once per facility per month.')
    para(doc, 'The EPV period is always the previous calendar month. Emails go to all four facility addresses on file (primary, owner, accounts, manager). Each email carries a personal token link into the EPV form.')

    h2(doc, 'Stage 2 — Submission')
    para(doc, 'The facility opens the token link. Company Admin or User completes the 4-step wizard (Business, Eggs, Pulp, Powder) and clicks Submit on Review. Purchase evidence is required for any Graded, Ungraded, Pulp, or Powder purchase greater than zero.')
    para(doc, 'On submit, the EPV is stamped Completed with the submitter\'s identity. The wizard becomes read-only.')

    h2(doc, 'Stage 3 — Inspector review')
    para(doc, 'The inspector assigned to the facility reviews the EPV. They either approve (facility figures stand) or reject and capture their own revised set as an Inspector EPV.')
    para(doc, 'The Inspector EPV, when present, is the source of truth for invoicing.')

    h2(doc, 'Stage 4 — ALS invoicing')
    para(doc, 'Once inspection-cleared, the EPV appears in the ALS Collections worklist. ALS marks the invoice as sent, uploads their invoice PDF, and reconciles the payment on the shared reconciliation fields. Every action is logged for transparency.')

    h2(doc, 'Stage 5 — Payment')
    para(doc, 'The facility uploads the Proof of Payment. ALS confirms and marks reconciled. The EPV is now closed for that month.')
    page_break(doc)

    h1(doc, 'Roles and access')
    section_intro(doc, 'Six roles exist in the system.')
    reference_table(doc,
        headers=['Role', 'Scope'],
        rows=[
            ['Super Admin',     'Full access to every page and every action.'],
            ['Admin',           'Dashboard, Clients, Company Overview, Administrators, Inspectors, User Management, Support.'],
            ['ALS',             'ALS Collections page, Company Overview (read-only cross-facility), Support, own Settings.'],
            ['Inspector',       'Inspectors dashboard, Company Overview, Support, own Settings.'],
            ['Company Admin',   'A single facility only (Company Overview, EPV form, Support, Settings).'],
            ['User',            'Subset of Company Admin (cannot invite others).'],
        ])
    page_break(doc)

    h1(doc, 'Audit and transparency')
    section_intro(doc, 'What is logged, where.')
    reference_table(doc,
        headers=['Log', 'Records'],
        rows=[
            ['ClientAuditLog',   'Every field-level change on a facility.'],
            ['EPVAuditLog',      'Every field-level change on an EPV, including ALS actions with the actor\'s role for the shared transparency log.'],
            ['LoginLog',         'Every login attempt — successful or failed — with reason, IP, and user agent.'],
            ['EmailSendLog',     'Every outbound email attempt, per recipient, with delivery status.'],
        ])
    page_break(doc)

    levy_summary(doc, technical=True)

    doc.save(path)


# ══════════════════════════════════════════════════════════════════════
# ENTRYPOINT
# ══════════════════════════════════════════════════════════════════════
BUILDERS = [
    ('EPVS Super Admin User Manual.docx',      build_super_admin),
    ('EPVS Admin User Manual.docx',            build_admin),
    ('EPVS Inspector User Manual.docx',        build_inspector),
    ('EPVS ALS User Manual.docx',              build_als),
    ('EPVS Company Admin User Manual.docx',    build_company_admin),
    ('EPVS User Manual.docx',                  build_user),
    ('EPVS Process Flow Documentation.docx',   build_process_flow),
]


def main():
    for name, builder in BUILDERS:
        target = os.path.join(REPO, name)
        try:
            builder(target)
            size_kb = os.path.getsize(target) // 1024
            print(f'wrote {name} ({size_kb} KB)')
        except Exception as e:
            print(f'FAILED {name}: {e}')
            raise


if __name__ == '__main__':
    main()
