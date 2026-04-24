"""
Rewrite all six EPVS user manuals from scratch, reflecting the current state
of the application (April 2026).

Each manual is produced as a fresh .docx with a cover page, table of
contents header, body sections and a support footer. Shared content blocks
are defined once; role-specific content lives in the build_*() functions.

Run:  python rewrite-manuals.py
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

REPO = os.path.dirname(os.path.abspath(__file__))
VERSION_LINE = f"Version: April 2026 rollout - last updated {date.today().isoformat()}"

BRAND_BLUE = RGBColor(0x14, 0x4E, 0x7A)
BRAND_GREY = RGBColor(0x55, 0x55, 0x55)


def set_run(run, size=None, bold=None, color=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def add_cover(doc, role_title):
    """Cover page: title, role, version."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(5):
        p.add_run().add_break()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p2.add_run("Egg Production Verification System"), size=24, bold=True, color=BRAND_BLUE)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p3.add_run(role_title), size=20, bold=True, color=BRAND_GREY)
    for _ in range(2):
        doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p4.add_run("ALS Food Safety Agency"), size=12)
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p5.add_run(VERSION_LINE), size=10, color=BRAND_GREY)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def H(doc, text, level=1):
    """Add a styled heading using a built-in style if available."""
    style_name = f'Heading {level}'
    try:
        doc.styles[style_name]
        return doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
        run.font.color.rgb = BRAND_BLUE
        return p


def P(doc, text):
    return doc.add_paragraph(text)


def B(doc, text):
    """Bulleted paragraph with safe fallback."""
    try:
        doc.styles['List Bullet']
        return doc.add_paragraph(text, style='List Bullet')
    except KeyError:
        p = doc.add_paragraph()
        p.add_run('  - ' + text)
        return p


def N(doc, items):
    """Numbered list, safe fallback."""
    try:
        doc.styles['List Number']
        for t in items:
            doc.add_paragraph(t, style='List Number')
        return
    except KeyError:
        for i, t in enumerate(items, 1):
            p = doc.add_paragraph()
            p.add_run(f'  {i}. {t}')


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# -----------------------------------------------------------------------------
# Shared content blocks
# -----------------------------------------------------------------------------

def intro_section(doc, role):
    H(doc, "1. Introduction")
    P(doc, f"This is the EPVS {role} user manual. It describes how to carry out "
           "every task your role requires inside the Egg Production Verification "
           "System (EPVS), a web application used by ALS Food Safety Agency "
           "to collect, verify, reconcile and report the monthly statutory "
           "egg and pulp levy returns submitted by production facilities.")

    H(doc, "1.1 System access", level=2)
    P(doc, "EPVS runs entirely in a web browser. Any recent version of Chrome, "
           "Edge or Firefox on a desktop, laptop or tablet is supported. The "
           "production URL is provided by your administrator. If you are asked "
           "to accept a security certificate or allow popups, say yes - "
           "certain workflows (for example opening a generated invoice or an "
           "uploaded Proof of Payment) rely on new tabs.")

    H(doc, "1.2 Who does what in EPVS", level=2)
    P(doc, "EPVS has five user roles. Each is scoped to a different part of "
           "the levy lifecycle:")
    B(doc, "Super Admin - organisation-wide oversight, user management, KPI "
           "targets, inspector approvals and full dashboards.")
    B(doc, "Admin - financial and reconciliation oversight; can manage "
           "companies and reconcile payments but cannot change system users or "
           "KPI targets.")
    B(doc, "Inspector - reviews facility EPVs against an independent "
           "inspector-captured EPV and approves or rejects them.")
    B(doc, "Company Admin - manages the users, data and EPV submissions for "
           "a single facility.")
    B(doc, "User - completes EPV submissions for their facility.")


def login_section(doc):
    H(doc, "2. Logging in")
    N(doc, [
        "Open the EPVS URL in your browser.",
        "Enter your email and password and click 'Sign In'.",
        "The first screen you see depends on your role - Super Admins land "
        "on the Super Admin Dashboard, Inspectors on the Inspector Dashboard, "
        "and company users on their Company Overview.",
    ])

    H(doc, "2.1 Forgotten password", level=2)
    N(doc, [
        "Click 'Forgot password?' on the login screen.",
        "Enter your account email.",
        "Open the password-reset email that arrives from the EPVS system "
        "mailbox.",
        "Click the link - it takes you to the Reset Password page.",
        "Enter and confirm a new password, then sign in.",
    ])
    P(doc, "Reset links expire after a short period. If yours has expired, "
           "request another from the same screen.")

    H(doc, "2.2 Accepting an invitation", level=2)
    P(doc, "New company users do not self-register. An administrator invites "
           "them via email. The invitation link opens the Accept Invite page "
           "where the invitee sets a password. Company Admins additionally "
           "step through a short facility wizard before their account is "
           "activated.")


def support_section(doc):
    H(doc, "Support tickets")
    P(doc, "A blue help button is pinned to the bottom right of every screen. "
           "Clicking it opens the Support window where you can either:")
    B(doc, "Browse your existing tickets to follow up on status or add a "
           "comment.")
    B(doc, "Click 'New ticket' to raise a new request. Choose a category - "
           "Administration questions route to the EPVS support team, while IT "
           "categories route to the technical team. Attach the ticket to a "
           "specific facility where it matters.")
    P(doc, "Every update on a ticket triggers an email. Tickets are closed "
           "once you confirm the matter is resolved.")


def epv_form_walkthrough(doc, audience):
    """Audience is 'facility' or 'inspector' - slightly different wording."""
    H(doc, "The EPV form explained")
    P(doc, "Every EPV is a four-step wizard that captures production, sales, "
           "deductions and pulp data for a given month and produces the "
           "statutory levy amount owed by the facility. EPV records are ALWAYS "
           "for a completed calendar month - an EPV received during March "
           "captures February's figures.")

    H(doc, "Step 1 - Business Details", level=2)
    P(doc, "Confirm or update the facility name, facility type, province, "
           "trading name, physical address, authorised person name and "
           "position, telephone numbers and email address. These fields are "
           "pre-filled from the Consolidated Master Facility Database but can "
           "be corrected here - the corrections flow back into the Company "
           "Overview audit log.")

    H(doc, "Step 2 - Eggs: Calculation of Statutory Levy", level=2)
    P(doc, "Capture everything in dozens. The form performs all totals "
           "automatically.")
    B(doc, "A. Opening Stock - carried forward from the previous month "
           "where possible, plus 'Eggs Produced during the Month'. Total A.")
    B(doc, "B. Purchases - Graded, Ungraded and Market Returns. When Graded "
           "or Ungraded is greater than zero, a 'Source & Supplier Detail' "
           "comment box and document uploader open; explain where and from "
           "whom the purchase was made, and attach PDFs, PNGs or JPGs of "
           "invoices or delivery notes up to 15 MB each. Total B.")
    B(doc, "C. Deductions - Machine Loss, Sent to Pulp, Destroyed and Eggs "
           "Exported. None of these trigger a levy but they reduce the "
           "theoretical closing stock. Total C.")
    B(doc, "D. Sales - Sold to Trade, Sold to Staff / Own Use, Sold through "
           "Farm Stall. Total D drives the egg levy: Total D x R0.020. This "
           "is shown immediately underneath the total.")
    B(doc, "E. Transfers - Sold or Transferred to Other Producers.")
    B(doc, "Closing Stock - the theoretical figure is computed as "
           "A + B - C - D - E. You then enter the 'Actual Closing Stock' "
           "(what you physically counted). If there is a difference, a "
           "variance reason field becomes mandatory.")

    H(doc, "Step 3 - Pulp: Calculation of Statutory Levy", level=2)
    P(doc, "All pulp quantities are captured in kilograms; the form "
           "automatically computes the dozen equivalent at 1.7 dozens per kg.")
    B(doc, "Opening Stock, Pulp Purchased from Others and Eggs Converted to "
           "Pulp combine to give Stock on Hand.")
    B(doc, "When 'Pulp Purchased from Others' is greater than zero, a "
           "separate comment box and uploader open for pulp purchase "
           "evidence. The 15 MB / PDF / PNG / JPG limits are identical to "
           "the egg page.")
    B(doc, "Sales to Trade is multiplied by the dozen conversion and by the "
           "levy rate to give the pulp levy amount.")
    B(doc, "Sold to Other Producers and Conversion Loss are subtracted from "
           "Stock on Hand to yield Closing Stock. Conversion Loss does NOT "
           "contribute to the pulp levy - it is purely a stock reconciliation "
           "entry.")

    H(doc, "Step 4 - Review & Submit", level=2)
    P(doc, "The review page is a read-only summary of everything you have "
           "entered, including uploaded documents and supplier detail. "
           "Check it carefully.")
    if audience == 'facility':
        P(doc, "Clicking 'Submit Verification' sends the completed EPV to "
               "the inspector and finance team. Once submitted, the facility "
               "cannot edit the EPV. Changes require a support ticket.")
    else:
        P(doc, "The inspector EPV is submitted the same way. It is compared "
               "against the facility's figures to produce the Verified / "
               "Rejected outcome on the Company Overview.")


def recent_updates_reference(doc):
    H(doc, "Appendix A - April 2026 change log")
    P(doc, "For users who knew the system before this release, the following "
           "items changed in April 2026:")
    B(doc, "EPVs are now always issued for the previous month (never the "
           "current month, which is not yet complete).")
    B(doc, "Facilities placed 'On EPV Cycle' now receive their monthly EPV "
           "automatically on the 1st of each month; no manual send is "
           "required after the first one.")
    B(doc, "The Company Overview now has a 'Resend Email' button on every "
           "Pending EPV to re-notify a facility that did not receive the "
           "original email.")
    B(doc, "The Consolidated Master Facility Database now shows a verified "
           "status badge per facility based on the outcome of their most "
           "recent EPV.")
    B(doc, "The EPV form adds 'Eggs Exported' as a deduction (does not "
           "affect the egg levy) and 'Conversion Loss' under the pulp page "
           "(does not affect the pulp levy, but reduces closing pulp stock).")
    B(doc, "Purchase evidence (supplier comment plus 15 MB PDF/PNG/JPG "
           "uploads) is available under both egg and pulp purchases.")
    B(doc, "The 'Add Verification' modal in Company Overview can only be "
           "used for past months - the current month is hidden and rejected "
           "by the backend.")


# -----------------------------------------------------------------------------
# Role-specific section builders
# -----------------------------------------------------------------------------

def super_admin_sections(doc):
    H(doc, "3. The Super Admin Dashboard")
    P(doc, "The Super Admin Dashboard is the system-wide health view. It "
           "opens every time you log in.")
    B(doc, "KPI tiles - Collection Rate, Approvals Actioned, Facilities "
           "Visited, Reconciliation Rate, Outstanding Rate and Verification "
           "Rate are each shown as a percentage against a configurable target. "
           "Click the pencil icon next to a tile to change its target; your "
           "change is logged and visible system-wide immediately.")
    B(doc, "Charts - time-series and breakdown charts show EPV submission "
           "volume, verification throughput, reconciliation progress and "
           "outstanding levies by month and province.")
    B(doc, "Action items - a scrollable list of tasks that need your "
           "attention: EPVs awaiting verification, overdue reconciliations, "
           "support tickets assigned to you.")

    H(doc, "4. Consolidated Master Facility Database")
    P(doc, "The master database is the single source of truth for every "
           "facility registered with EPVS.")
    B(doc, "Search by business name, account code, email, town or province "
           "via the toolbar. Results paginate 50 at a time.")
    B(doc, "Click 'Edit' on any row to change facility details - every "
           "change is written to the audit log with your name and the "
           "timestamp.")
    B(doc, "Click 'Hist' to open the audit history for just that facility.")
    B(doc, "The status column now shows the verification outcome of the "
           "facility's most recent EPV: No EPV, Pending Submission, Pending "
           "Verification or Verified. An 'On EPV Cycle' tag additionally "
           "indicates that the facility is on the monthly auto-send schedule.")
    B(doc, "Click 'Send EPV' to issue the first EPV for the previous month "
           "and place the facility on the monthly cycle. The email goes to "
           "the primary, owner, accounts and manager addresses on file; each "
           "recipient is tracked separately in EmailSendLog.")
    B(doc, "Use 'Add Client' to register a new facility. The 'Delete' button "
           "removes a facility and cascades through related EPVs, invitations "
           "and audit entries - handle with care.")

    H(doc, "5. Company Overview")
    P(doc, "Click a facility name anywhere in EPVS to open its Company "
           "Overview. This is the operational cockpit for a single facility.")
    B(doc, "Tab 1 - Facility Details: same fields as the master database, "
           "with the full audit log exposed in a right-hand drawer.")
    B(doc, "Tab 2 - Users: invited company users. Invite a new Company "
           "Admin or User, re-send an outstanding invitation, or deactivate "
           "an existing user.")
    B(doc, "Tab 3 - EPVs: every EPV issued for this facility, newest first. "
           "Each row shows reference number, period, completion date, and "
           "separate action columns for the facility EPV, the inspector EPV, "
           "the verification outcome, POP upload, reconciliation and manual "
           "inspection.")
    B(doc, "Pending EPVs expose a 'Resend Email' button - use it when the "
           "facility reports they never got the original.")
    B(doc, "Completed but unverified EPVs can be approved or rejected "
           "directly from this table by Super Admins and Inspectors.")

    H(doc, "6. User Management (Settings)")
    P(doc, "Open the Settings icon in the top-right to manage EPVS users "
           "(not company-level users - those live under Company Overview).")
    B(doc, "Search, filter by role and active/deactivated status.")
    B(doc, "Edit a user's first name, last name, email and role.")
    B(doc, "Deactivate a user - they retain their history but cannot log in. "
           "Reactivation is a one-click reverse.")
    B(doc, "Reset password - sends the user a password-reset email.")

    H(doc, "7. The Inspector and Administrator dashboards")
    P(doc, "Super Admins can view both the Inspector and Administrator "
           "dashboards via the navbar. These give you the same operational "
           "lens that inspectors and admins see, useful for spotting a stuck "
           "workflow or supporting a colleague.")


def admin_sections(doc):
    H(doc, "3. The Administrator Dashboard")
    P(doc, "The Administrator Dashboard focuses on financial throughput.")
    B(doc, "Reconciliation KPIs - reconciled vs outstanding levies per "
           "month.")
    B(doc, "Outstanding facilities - a queue of EPVs that have uploaded a "
           "Proof of Payment but have not yet been reconciled.")
    B(doc, "Batch reconcile - tick several EPVs and reconcile them in one "
           "action.")

    H(doc, "4. Consolidated Master Facility Database")
    P(doc, "Admins have the same view as Super Admins over the master "
           "facility list. You can edit facility details, issue EPVs, place "
           "facilities on the monthly cycle, and see the new verified status "
           "badge (No EPV / Pending Submission / Pending Verification / "
           "Verified).")

    H(doc, "5. Company Overview")
    P(doc, "Same cockpit view as Super Admins, with a few restrictions: you "
           "cannot change user roles or KPI targets, but you can reconcile "
           "POPs, resend EPV emails, mark EPVs as manually inspected and "
           "leave comments.")

    H(doc, "6. Reconciliation workflow")
    N(doc, [
        "Facility uploads Proof of Payment inside the Company Overview.",
        "Open the POP row and click the thumbnail to view.",
        "If the amount matches the levy, tick 'Reconciled'. Optionally "
        "capture a reconciled amount (for partial payments).",
        "Rejected POPs leave the EPV in its current state - leave an inline "
        "comment explaining why; this appears in the audit log and notifies "
        "the company admin.",
    ])


def inspector_sections(doc):
    H(doc, "3. The Inspector Dashboard")
    B(doc, "Pending Approvals - completed facility EPVs assigned to your "
           "province that need to be approved or rejected.")
    B(doc, "Visits - quarterly facility visits scheduled for you.")
    B(doc, "Reconciliation overview - the levy amounts linked to EPVs you "
           "have approved.")

    H(doc, "4. Verifying a facility EPV")
    N(doc, [
        "Click the reference number to open the facility's submitted EPV.",
        "Cross-check every figure against source documentation (purchase "
        "invoices, production logs).",
        "If the figures match, click 'Approve' on the verify column. The "
        "status flips to Verified and the facility is notified.",
        "If the figures do not match, capture your own figures via "
        "'Create Inspector EPV' on the same row. Submit the inspector EPV. "
        "The facility EPV is automatically marked Rejected if the two "
        "differ, and a comparison panel becomes visible on the Company "
        "Overview for Super Admins and Admins to resolve.",
    ])
    P(doc, "Inspector EPVs use the same four-step wizard as facility EPVs. "
           "The period, reference number and pre-fill data are all inherited "
           "from the linked facility EPV.")

    H(doc, "5. The new EPV form fields")
    P(doc, "When you open either a facility or an inspector EPV, be aware of "
           "the updated form fields from the April 2026 release:")
    B(doc, "Section C (Deductions) now includes 'Eggs Exported'. Export "
           "volumes should be removed here, not in Sales.")
    B(doc, "Under Section B and under pulp Stock In, any non-zero purchase "
           "now triggers a comment box and document uploader where the "
           "facility should name the supplier and attach an invoice or "
           "delivery note. If these are missing, reject the EPV and ask for "
           "them.")
    B(doc, "The pulp page carries a separate 'Conversion Loss' line - this "
           "affects pulp closing stock but not the pulp levy.")


def company_admin_sections(doc):
    H(doc, "3. First-time setup")
    N(doc, [
        "Open the invitation email from 'automaticmails@...' and click "
        "'Accept Invitation'.",
        "Set your password.",
        "Step through the facility wizard, confirming each pre-filled "
        "detail. Anything you change here is committed to the facility "
        "record and logged.",
        "Click 'Finish'. EPVS drops you on your Company Overview page.",
    ])

    H(doc, "4. The Company Overview")
    B(doc, "Tab 1 - Facility Details: edit your facility profile. Any "
           "change is recorded in the audit drawer.")
    B(doc, "Tab 2 - Users: invite basic users at your facility, re-send "
           "outstanding invitations, or deactivate users who no longer work "
           "with you.")
    B(doc, "Tab 3 - EPVs: your month-by-month verification list. Open, "
           "complete and track the status of each.")

    H(doc, "5. Completing an EPV")
    P(doc, "EPVs arrive in your email inbox once a month. They always "
           "capture the PREVIOUS calendar month's data, not the current "
           "month. You can also self-open an overdue EPV from Tab 3 of your "
           "Company Overview.")
    P(doc, "The form is a four-step wizard - see the 'EPV form explained' "
           "chapter for the detail on each step. Two operational tips:")
    B(doc, "If you purchased eggs or pulp from another facility during the "
           "month, be sure to name the supplier in the Source & Supplier "
           "Detail box that opens under each purchase section, and attach "
           "the invoice as a PDF or photo. Without this, inspectors will "
           "reject the EPV.")
    B(doc, "If the 'Actual Closing Stock' you physically counted differs "
           "from the theoretical figure, you must explain the variance in "
           "the box that appears. Typical reasons: tray breakage, on-farm "
           "consumption, stock count timing differences.")

    H(doc, "6. Adding a verification for a previous month")
    P(doc, "If you need to submit an EPV for a past period that was never "
           "issued, open Company Overview, click 'Add Verification' above "
           "the EPV list, choose the month (only completed months appear) "
           "and year, and click 'Create & Complete'. The form opens in the "
           "same wizard. You cannot add a verification for the current "
           "month - the system only accepts months that have ended.")

    H(doc, "7. Uploading Proof of Payment")
    N(doc, [
        "In Tab 3, click the POP cell for the row you have paid.",
        "Drag a PDF or image file onto the uploader, or click to browse. "
        "Max 10 MB per POP file.",
        "The admin team receives a notification and reconciles the payment. "
        "You can see the reconciled status update live in the same row.",
    ])

    H(doc, "8. If you did not receive your EPV email")
    P(doc, "Ask your Super Admin or the EPVS support team to resend. On the "
           "Company Overview, next to every Pending EPV, there is now a "
           "'Resend Email' button that re-sends the original notification "
           "to every email address held on file for your facility. No "
           "duplicate EPV is created.")


def user_sections(doc):
    H(doc, "3. Your role")
    P(doc, "Basic users submit EPV forms for their facility. Invitations to "
           "other company users, facility details and POP uploads are handled "
           "by your Company Admin.")

    H(doc, "4. Completing an EPV")
    P(doc, "When the monthly EPV notification arrives in your inbox, click "
           "the big 'Complete Verification' button in the email. The four-"
           "step wizard opens - see the 'EPV form explained' chapter.")
    P(doc, "Two frequently asked questions:")
    B(doc, "\"I cannot find the email.\" Contact your Company Admin - they "
           "have a 'Resend Email' button inside Company Overview.")
    B(doc, "\"The month is wrong.\" The EPV always captures the previous "
           "calendar month. In March, you are submitting February figures. "
           "This is deliberate - the current month is not yet complete.")

    H(doc, "5. The purchase source and document uploads")
    P(doc, "If your facility bought eggs or pulp from another producer during "
           "the month, a comment box opens under that section. Describe the "
           "supplier briefly (name and any reference). Attach a scanned "
           "invoice or delivery note if you have it - PDF, PNG or JPG up to "
           "15 MB. Without this, your EPV may be rejected by the inspector.")


def process_flow_sections(doc):
    H(doc, "1. Document purpose")
    P(doc, "This document walks through the end-to-end lifecycle of a single "
           "statutory levy return inside EPVS, from the first time a facility "
           "is placed on the EPV cycle through to levy reconciliation. It is "
           "aimed at anyone implementing or auditing the process.")

    H(doc, "2. Actor roles")
    B(doc, "Super Admin - operates EPVS; has system-wide access.")
    B(doc, "Admin - operates EPVS; focuses on financial reconciliation.")
    B(doc, "Inspector - field role; verifies facility EPVs against an "
           "independently captured inspector EPV.")
    B(doc, "Company Admin - manages a single facility inside EPVS.")
    B(doc, "User - captures EPVs for a single facility.")

    H(doc, "3. The monthly cycle")
    N(doc, [
        "On the 1st of each month, the EPVS scheduler iterates every "
        "facility with EPVCycleStatus = 'On EPV Cycle' and issues an EPV "
        "for the previous calendar month (so an EPV issued on 1 April 2026 "
        "captures March 2026 figures).",
        "The EPV email is sent to every valid address on file for that "
        "facility (primary, owner, accounts, manager). Each send is logged "
        "separately in EmailSendLog so failures are visible.",
        "The facility opens the email, completes the wizard and submits. "
        "Submission is final for the facility - further changes require a "
        "support ticket.",
        "The Inspector for that province reviews the EPV. They either "
        "Approve it outright or capture a parallel Inspector EPV with their "
        "own figures. If the Inspector EPV differs from the facility EPV, "
        "the facility EPV is marked Rejected and surfaced to Super Admin / "
        "Admin for resolution.",
        "Once verified, the facility uploads Proof of Payment. An Admin "
        "reconciles the levy against the POP.",
        "The cycle repeats for the next month.",
    ])

    H(doc, "4. Putting a new facility on the cycle")
    P(doc, "A facility is placed on the monthly cycle by issuing its first "
           "EPV from either the Consolidated Master Facility Database or "
           "from Company Overview. The 'Send EPV' action sets "
           "EPVCycleStatus = 'On EPV Cycle' on the facility record. From "
           "then on the scheduler takes over.")
    P(doc, "To take a facility off the cycle, clear EPVCycleStatus from "
           "their facility record.")

    H(doc, "5. Data captured per EPV")
    P(doc, "Each EPV stores two parallel sets of figures - the facility's "
           "own submission and the inspector's parallel capture. Fields are "
           "grouped into:")
    B(doc, "Business details (name, type, province, authorised person, "
           "contact channels).")
    B(doc, "A - Opening stock and production.")
    B(doc, "B - Purchases with supplier detail and attached invoices "
           "(up to 15 MB per file, PDF / PNG / JPG).")
    B(doc, "C - Deductions (Machine Loss, Sent to Pulp, Destroyed, Eggs "
           "Exported).")
    B(doc, "D - Sales, which drive the egg levy at R0.020 per dozen.")
    B(doc, "E - Transfers to other producers.")
    B(doc, "Pulp opening stock, purchases with supplier detail, "
           "conversion from eggs, sales to trade (drives pulp levy), "
           "transfers to producers, and conversion loss.")
    B(doc, "Closing stock with variance explanation when actual differs "
           "from theoretical.")

    H(doc, "6. Reconciliation workflow")
    N(doc, [
        "Facility uploads POP.",
        "Admin reviews, optionally captures a reconciled amount, flags "
        "'Reconciled'.",
        "If a POP is disputed, admin leaves an inline comment - the "
        "facility is notified by email.",
    ])

    H(doc, "7. Support and audit")
    P(doc, "Every meaningful change (facility edit, EPV submission, "
           "verification, POP upload, reconciliation, resend) is written "
           "to either ClientAuditLog (company scope) or EPVAuditLog (EPV "
           "scope). Support tickets are the official change channel for "
           "completed EPVs and link back to the facility.")


# -----------------------------------------------------------------------------
# Manual assembly
# -----------------------------------------------------------------------------

def build(filename, role_title, role_sections, audience_for_form=None,
          include_form_walkthrough=True, include_support=True):
    doc = Document()
    add_cover(doc, role_title)
    intro_section(doc, role_title.replace(' User Manual', '').replace('EPVS ', ''))
    page_break(doc)
    login_section(doc)
    page_break(doc)
    role_sections(doc)
    if include_form_walkthrough:
        page_break(doc)
        epv_form_walkthrough(doc, audience_for_form or 'facility')
    if include_support:
        page_break(doc)
        support_section(doc)
    page_break(doc)
    recent_updates_reference(doc)
    doc.save(filename)


def main():
    targets = [
        ("EPVS Super Admin User Manual.docx", "Super Admin User Manual",
            super_admin_sections, 'inspector', True, True),
        ("EPVS Admin User Manual.docx", "Administrator User Manual",
            admin_sections, 'facility', True, True),
        ("EPVS Inspector User Manual.docx", "Inspector User Manual",
            inspector_sections, 'inspector', True, True),
        ("EPVS Company Admin User Manual.docx", "Company Admin User Manual",
            company_admin_sections, 'facility', True, True),
        ("EPVS User Manual.docx", "User Manual",
            user_sections, 'facility', True, True),
    ]

    for name, title, builder, audience, include_form, include_support in targets:
        path = os.path.join(REPO, name)
        build(path, title, builder, audience, include_form, include_support)
        print("wrote:", name)

    # Process Flow is structured differently - no login/form walkthrough.
    pf_path = os.path.join(REPO, "EPVS Process Flow Documentation.docx")
    doc = Document()
    add_cover(doc, "Process Flow Documentation")
    process_flow_sections(doc)
    page_break(doc)
    recent_updates_reference(doc)
    doc.save(pf_path)
    print("wrote: EPVS Process Flow Documentation.docx")


if __name__ == "__main__":
    main()
